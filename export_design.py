import json
from docx import Document

with open("cedehub_design.json", "r", encoding="utf-8") as f:
    data = json.load(f)

messages = data.get("messages", {})
colors = data.get("colors", [])
fonts = data.get("fonts", [])
tokens = data.get("design_tokens", {})

# =====================
# Xuất file .md
# =====================
md = []
md.append("# Cedehub.io — Design Document\n")

md.append("## Message & Content\n")
md.append(f"**Hero Title:** {messages.get('hero_title', '')}\n")

md.append("\n### Headings")
for h in messages.get("headings", []):
    md.append(f"- {h}")

md.append("\n### Descriptions")
for d in messages.get("descriptions", []):
    md.append(f"- {d}")

md.append("\n### Buttons / CTAs")
for b in messages.get("buttons", []):
    md.append(f"- {b}")

md.append("\n## Color Palette\n")
for c in colors:
    md.append(f"- `{c}`")

md.append("\n## Typography / Fonts\n")
for f in fonts:
    md.append(f"- {f}")

md.append("\n## Design Tokens\n")
for k, v in tokens.items():
    md.append(f"- `{k}`: {v}")

with open("cedehub_design.md", "w", encoding="utf-8") as f:
    f.write("\n".join(md))

print("✅ Đã xuất cedehub_design.md")

# =====================
# Xuất file .docx
# =====================
doc = Document()
doc.add_heading("Cedehub.io — Design Document", 0)

doc.add_heading("Message & Content", 1)
doc.add_paragraph(f"Hero Title: {messages.get('hero_title', '')}")

doc.add_heading("Headings", 2)
for h in messages.get("headings", []):
    doc.add_paragraph(h, style="List Bullet")

doc.add_heading("Descriptions", 2)
for d in messages.get("descriptions", []):
    doc.add_paragraph(d, style="List Bullet")

doc.add_heading("Buttons / CTAs", 2)
for b in messages.get("buttons", []):
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Color Palette", 1)
for c in colors:
    doc.add_paragraph(c, style="List Bullet")

doc.add_heading("Typography / Fonts", 1)
for f in fonts:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("Design Tokens", 1)
for k, v in tokens.items():
    doc.add_paragraph(f"{k}: {v}", style="List Bullet")

doc.save("cedehub_design.docx")
print("✅ Đã xuất cedehub_design.docx")